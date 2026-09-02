from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any, Protocol

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.exceptions import PackageNotFoundError
from docx.text.paragraph import Paragraph

from markdown_docx.assets import default_template_bytes
from markdown_docx.errors import TemplateError


class Story(Protocol):
    @property
    def paragraphs(self) -> list[Paragraph]: ...

    @property
    def tables(self) -> list[Any]: ...


def load_template(path: Path | None) -> DocumentObject:
    if path is None:
        document = Document(BytesIO(default_template_bytes()))
        validate_blank_template(document, label="packaged default template")
        return document
    resolved = path.resolve()
    if not resolved.is_file():
        raise TemplateError("template_not_found", f"Template does not exist: {resolved}")
    if resolved.suffix.lower() != ".docx":
        raise TemplateError("unsupported_feature", "Templates must use the .docx format. DOTX is not supported.")
    try:
        document = Document(str(resolved))
    except (PackageNotFoundError, ValueError) as exc:
        raise TemplateError("template_invalid", f"Template is not a readable DOCX file: {resolved}") from exc
    validate_blank_template(document, label=str(resolved))
    return document


def validate_blank_template(document: DocumentObject, *, label: str) -> None:
    errors = blank_template_errors(document)
    if errors:
        raise TemplateError(
            "template_not_blank",
            f"Template must be blank and formatting-only: {errors[0]}",
            details={"template": label, "errors": errors},
        )


def blank_template_errors(document: DocumentObject) -> list[str]:
    errors: list[str] = []
    if any(paragraph.text.strip() for paragraph in document.paragraphs):
        errors.append("body contains text")
    if document.tables:
        errors.append("body contains tables")
    if len(document.inline_shapes) or any(_paragraph_has_drawing(p) for p in document.paragraphs):
        errors.append("body contains images or drawings")
    for number, section in enumerate(document.sections, start=1):
        for story_name, story in (("header", section.header), ("footer", section.footer)):
            if _story_has_content(story):
                errors.append(f"section {number} {story_name} is not empty")
    return errors


def inspect_template(path: Path | None) -> dict[str, Any]:
    if path is None:
        label = "packaged-default"
        document = Document(BytesIO(default_template_bytes()))
    else:
        resolved = path.resolve()
        label = str(resolved)
        if not resolved.is_file():
            raise TemplateError("template_not_found", f"Template does not exist: {resolved}")
        try:
            document = Document(str(resolved))
        except (PackageNotFoundError, ValueError) as exc:
            raise TemplateError("template_invalid", f"Template is not a readable DOCX file: {resolved}") from exc
    errors = blank_template_errors(document)
    groups = {
        "paragraph": sorted(style.name for style in document.styles if style.type == WD_STYLE_TYPE.PARAGRAPH),
        "character": sorted(style.name for style in document.styles if style.type == WD_STYLE_TYPE.CHARACTER),
        "table": sorted(style.name for style in document.styles if style.type == WD_STYLE_TYPE.TABLE),
    }
    sections = [
        {
            "index": index,
            "orientation": section.orientation.name.lower() if section.orientation else None,
            "width_inches": round(section.page_width.inches, 4) if section.page_width else None,
            "height_inches": round(section.page_height.inches, 4) if section.page_height else None,
            "margins_inches": {
                "top": round(section.top_margin.inches, 4) if section.top_margin else None,
                "right": round(section.right_margin.inches, 4) if section.right_margin else None,
                "bottom": round(section.bottom_margin.inches, 4) if section.bottom_margin else None,
                "left": round(section.left_margin.inches, 4) if section.left_margin else None,
            },
        }
        for index, section in enumerate(document.sections, start=1)
    ]
    return {
        "template": label,
        "valid": not errors,
        "errors": errors,
        "styles": groups,
        "sections": sections,
    }


def _story_has_content(story: Story) -> bool:
    return bool(story.tables) or any(
        paragraph.text.strip() or _paragraph_has_drawing(paragraph) for paragraph in story.paragraphs
    )


def _paragraph_has_drawing(paragraph: Paragraph) -> bool:
    for run in paragraph.runs:
        if any(not isinstance(item, str) for item in run.iter_inner_content()):
            return True
    return False
