from __future__ import annotations

import os
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Emu
from docx.text.paragraph import Paragraph

from markdown_docx.errors import MarkdownDocxError, RenderError
from markdown_docx.images import ImageLoader, rendered_width
from markdown_docx.models import (
    CodeBlock,
    DocumentModel,
    HeadingBlock,
    ImageBlock,
    InlineFragment,
    ListParagraphBlock,
    PageBreakBlock,
    ParagraphBlock,
    SectionBreakBlock,
    SectionSettings,
    TableBlock,
)
from markdown_docx.styles import apply_font_overrides, validate_styles
from markdown_docx.template import load_template

PARAGRAPH_ALIGNMENT = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}
TABLE_ALIGNMENT = {
    "left": WD_TABLE_ALIGNMENT.LEFT,
    "center": WD_TABLE_ALIGNMENT.CENTER,
    "right": WD_TABLE_ALIGNMENT.RIGHT,
}


def render_docx(
    model: DocumentModel,
    output_path: Path,
    *,
    template_path: Path | None,
    base_dir: Path,
    allow_remote_images: bool,
) -> dict[str, Any]:
    document = load_template(template_path)
    validate_styles(document, model.options)
    apply_font_overrides(document, model.options)
    current_settings = model.options.section
    _apply_section_settings(document.sections[0], current_settings)
    image_loader = ImageLoader(base_dir, allow_remote=allow_remote_images)
    reusable = _reusable_initial_paragraph(document)
    warnings = list(model.warnings)
    image_seen = False

    try:
        for block in model.blocks:
            if isinstance(block, PageBreakBlock):
                document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            elif isinstance(block, SectionBreakBlock):
                current_settings = block.settings
                section = document.add_section(WD_SECTION.NEW_PAGE)
                _apply_section_settings(section, current_settings)
            elif isinstance(block, HeadingBlock):
                paragraph, reusable = _new_paragraph(
                    document,
                    style=model.options.styles.headings[block.level],
                    reusable=reusable,
                )
                image_seen |= _render_fragments(
                    paragraph,
                    block.fragments,
                    image_loader=image_loader,
                    settings=current_settings,
                    monospace=model.options.fonts.monospace,
                    line=block.line,
                    input_path=model.source_name,
                )
            elif isinstance(block, ParagraphBlock):
                style = (
                    model.options.styles.blockquote if block.role == "blockquote" else model.options.styles.paragraph
                )
                paragraph, reusable = _new_paragraph(document, style=style, reusable=reusable)
                image_seen |= _render_fragments(
                    paragraph,
                    block.fragments,
                    image_loader=image_loader,
                    settings=current_settings,
                    monospace=model.options.fonts.monospace,
                    line=block.line,
                    input_path=model.source_name,
                )
            elif isinstance(block, CodeBlock):
                paragraph, reusable = _new_paragraph(
                    document,
                    style=model.options.styles.code_block,
                    reusable=reusable,
                )
                paragraph.add_run(block.text.rstrip("\n"))
            elif isinstance(block, ListParagraphBlock):
                styles = (
                    model.options.styles.ordered_list
                    if block.list_kind == "ordered"
                    else model.options.styles.unordered_list
                )
                paragraph, reusable = _new_paragraph(document, style=styles[block.depth], reusable=reusable)
                image_seen |= _render_fragments(
                    paragraph,
                    block.fragments,
                    image_loader=image_loader,
                    settings=current_settings,
                    monospace=model.options.fonts.monospace,
                    line=block.line,
                    input_path=model.source_name,
                )
            elif isinstance(block, TableBlock):
                _render_table(
                    document,
                    block,
                    model=model,
                    settings=current_settings,
                    image_loader=image_loader,
                )
            elif isinstance(block, ImageBlock):
                paragraph, reusable = _new_paragraph(
                    document,
                    style=model.options.styles.paragraph,
                    reusable=reusable,
                )
                paragraph.alignment = PARAGRAPH_ALIGNMENT[block.options.alignment]
                asset = image_loader.load(block.src, line=block.line, input_path=model.source_name)
                width = rendered_width(
                    asset,
                    block.options,
                    usable_width=current_settings.usable_width,
                    line=block.line,
                    input_path=model.source_name,
                )
                paragraph.add_run().add_picture(BytesIO(asset.data), width=Emu(width))
                image_seen = True
            else:
                raise AssertionError(f"Unhandled block type: {type(block).__name__}")
        if image_seen:
            warnings.append("image_alt_text_not_embedded")
        _save_atomically(document, output_path)
    except MarkdownDocxError:
        raise
    except Exception as exc:
        raise RenderError("render_failed", f"Could not render DOCX: {exc}") from exc

    return {
        "output": str(output_path),
        "sections": len(document.sections),
        "warnings": sorted(set(warnings)),
    }


def _apply_section_settings(section: Any, settings: SectionSettings) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE if settings.orientation == "landscape" else WD_ORIENT.PORTRAIT
    section.page_width = Emu(settings.effective_width)
    section.page_height = Emu(settings.effective_height)
    section.top_margin = Emu(settings.margins.top)
    section.right_margin = Emu(settings.margins.right)
    section.bottom_margin = Emu(settings.margins.bottom)
    section.left_margin = Emu(settings.margins.left)


def _reusable_initial_paragraph(document: DocumentObject) -> Paragraph | None:
    if len(document.paragraphs) == 1 and not document.tables:
        paragraph = document.paragraphs[0]
        if not paragraph.text and not paragraph.runs:
            return paragraph
    return None


def _new_paragraph(
    document: DocumentObject,
    *,
    style: str,
    reusable: Paragraph | None,
) -> tuple[Paragraph, None]:
    if reusable is not None:
        reusable.style = style
        return reusable, None
    return document.add_paragraph(style=style), None


def _render_fragments(
    paragraph: Paragraph,
    fragments: list[InlineFragment],
    *,
    image_loader: ImageLoader,
    settings: SectionSettings,
    monospace: str,
    line: int,
    input_path: str,
) -> bool:
    image_seen = False
    for fragment in fragments:
        if fragment.kind == "break":
            paragraph.add_run().add_break(WD_BREAK.LINE)
        elif fragment.kind == "image":
            asset = image_loader.load(fragment.src or "", line=line, input_path=input_path)
            width = min(asset.natural_width, settings.usable_width)
            run = paragraph.add_run()
            run.bold = fragment.bold or None
            run.italic = fragment.italic or None
            run.add_picture(BytesIO(asset.data), width=Emu(width))
            image_seen = True
        else:
            run = paragraph.add_run(fragment.text or "")
            run.bold = fragment.bold or None
            run.italic = fragment.italic or None
            if fragment.code:
                run.font.name = monospace
    return image_seen


def _render_table(
    document: DocumentObject,
    block: TableBlock,
    *,
    model: DocumentModel,
    settings: SectionSettings,
    image_loader: ImageLoader,
) -> None:
    row_data = [block.headers, *block.rows]
    column_count = len(block.headers)
    style_name = block.options.style or model.options.styles.table
    table = document.add_table(rows=len(row_data), cols=column_count, style=style_name)
    table.alignment = TABLE_ALIGNMENT[block.options.alignment]
    set_widths = block.options.width == "page" or block.options.column_widths is not None
    table.autofit = not set_widths
    widths: list[int] = []
    if set_widths:
        ratios = list(block.options.column_widths or (1.0,) * column_count)
        ratio_total = sum(ratios)
        widths = [round(settings.usable_width * ratio / ratio_total) for ratio in ratios]
        widths[-1] += settings.usable_width - sum(widths)
        for column, width in zip(table.columns, widths, strict=True):
            column.width = Emu(width)

    for row_index, cells in enumerate(row_data):
        for column_index, cell_data in enumerate(cells):
            cell = table.cell(row_index, column_index)
            if widths:
                cell.width = Emu(widths[column_index])
            paragraph = cell.paragraphs[0]
            paragraph.style = model.options.styles.paragraph
            paragraph.alignment = PARAGRAPH_ALIGNMENT[cell_data.alignment]
            _render_fragments(
                paragraph,
                cell_data.fragments,
                image_loader=image_loader,
                settings=settings,
                monospace=model.options.fonts.monospace,
                line=block.line,
                input_path=model.source_name,
            )


def _save_atomically(document: DocumentObject, output_path: Path) -> None:
    output_path = output_path.resolve()
    if not output_path.parent.is_dir():
        raise RenderError("render_failed", f"Output directory does not exist: {output_path.parent}")
    descriptor, temporary_name = tempfile.mkstemp(
        prefix=f".{output_path.stem}-", suffix=".docx", dir=output_path.parent
    )
    os.close(descriptor)
    temporary_path = Path(temporary_name)
    try:
        document.save(str(temporary_path))
        Document(str(temporary_path))
        temporary_path.replace(output_path)
    finally:
        temporary_path.unlink(missing_ok=True)
