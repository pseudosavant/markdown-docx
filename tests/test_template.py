from __future__ import annotations

from collections.abc import Callable
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from markdown_docx.errors import TemplateError
from markdown_docx.parser import parse_document
from markdown_docx.renderer import render_docx
from markdown_docx.template import inspect_template, load_template


def test_packaged_default_template_is_blank_and_has_required_styles() -> None:
    details = inspect_template(None)
    assert details["valid"] is True
    assert "Code Block" in details["styles"]["paragraph"]
    assert "Table Grid" in details["styles"]["table"]


def test_blank_custom_template_is_valid(blank_template_factory: Callable[[str], Path]) -> None:
    path = blank_template_factory("blank.docx")
    assert inspect_template(path)["valid"] is True
    load_template(path)


@pytest.mark.parametrize("content_kind", ["text", "table", "image", "header", "footer"])
def test_nonblank_custom_templates_are_rejected(tmp_path: Path, content_kind: str, png_file: Path) -> None:
    document = Document()
    if "Code Block" not in document.styles:
        document.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    if content_kind == "text":
        document.add_paragraph("Existing content")
    elif content_kind == "table":
        document.add_table(rows=1, cols=1)
    elif content_kind == "image":
        document.add_picture(str(png_file))
    elif content_kind == "header":
        document.sections[0].header.paragraphs[0].text = "Header"
    else:
        document.sections[0].footer.paragraphs[0].text = "Footer"
    path = tmp_path / f"{content_kind}.docx"
    document.save(path)
    with pytest.raises(TemplateError) as excinfo:
        load_template(path)
    assert excinfo.value.context.code == "template_not_blank"


def test_missing_and_wrong_type_styles_are_rejected(tmp_path: Path) -> None:
    missing = tmp_path / "missing.docx"
    Document().save(missing)
    model = parse_document("Text\n", input_path=tmp_path / "input.md", source_name="input.md")
    with pytest.raises(TemplateError) as missing_error:
        render_docx(model, tmp_path / "out.docx", template_path=missing, base_dir=tmp_path, allow_remote_images=False)
    assert missing_error.value.context.code == "template_style_missing"

    wrong_document = Document()
    wrong_document.styles.add_style("Code Block", WD_STYLE_TYPE.CHARACTER)
    wrong = tmp_path / "wrong.docx"
    wrong_document.save(wrong)
    with pytest.raises(TemplateError) as wrong_error:
        render_docx(model, tmp_path / "out.docx", template_path=wrong, base_dir=tmp_path, allow_remote_images=False)
    assert wrong_error.value.context.code == "template_style_type_mismatch"


def test_template_theme_and_numbering_parts_are_preserved(
    tmp_path: Path,
    blank_template_factory: Callable[[str], Path],
) -> None:
    template = blank_template_factory("formatting.docx")
    before = template.read_bytes()
    model = parse_document("1. Item\n", input_path=tmp_path / "input.md", source_name="input.md")
    output = tmp_path / "output.docx"
    render_docx(model, output, template_path=template, base_dir=tmp_path, allow_remote_images=False)
    assert template.read_bytes() == before
    with ZipFile(template) as source, ZipFile(output) as rendered:
        assert rendered.read("word/theme/theme1.xml") == source.read("word/theme/theme1.xml")
        assert rendered.read("word/numbering.xml") == source.read("word/numbering.xml")


def test_dotx_is_rejected(tmp_path: Path) -> None:
    path = tmp_path / "template.dotx"
    path.write_bytes(b"not a package")
    with pytest.raises(TemplateError) as excinfo:
        load_template(path)
    assert excinfo.value.context.code == "unsupported_feature"
