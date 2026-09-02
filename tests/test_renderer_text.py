from __future__ import annotations

from pathlib import Path

from docx import Document

from markdown_docx.parser import parse_document
from markdown_docx.renderer import render_docx


def render(tmp_path: Path, source: str) -> Path:
    output = tmp_path / "output.docx"
    model = parse_document(source, input_path=tmp_path / "input.md", source_name=str(tmp_path / "input.md"))
    render_docx(model, output, template_path=None, base_dir=tmp_path, allow_remote_images=False)
    return output


def test_render_text_styles_and_inline_formatting(tmp_path: Path) -> None:
    output = render(
        tmp_path,
        """# Heading

Plain **bold** and *italic* with `code`.\\
Second line.

> Quoted text.

```python
print("hello")
print("world")
```
""",
    )
    document = Document(output)
    assert [paragraph.style.name for paragraph in document.paragraphs] == [
        "Heading 1",
        "Normal",
        "Quote",
        "Code Block",
    ]
    paragraph = document.paragraphs[1]
    assert any(run.text == "bold" and run.bold for run in paragraph.runs)
    assert any(run.text == "italic" and run.italic for run in paragraph.runs)
    assert any(run.text == "code" and run.font.name == "Consolas" for run in paragraph.runs)
    assert "Second line" in paragraph.text
    assert document.paragraphs[3].text == 'print("hello")\nprint("world")'


def test_document_font_overrides_update_mapped_styles(tmp_path: Path) -> None:
    output = render(
        tmp_path,
        """<!-- markdown-docx
document:
  fonts:
    body: Arial
    headings: Arial
    monospace: Cascadia Mono
-->

# Heading

Body with `code`.
""",
    )
    document = Document(output)
    assert document.styles["Normal"].font.name == "Arial"
    assert document.styles["Heading 1"].font.name == "Arial"
    assert document.styles["Code Block"].font.name == "Cascadia Mono"
    assert next(run for run in document.paragraphs[1].runs if run.text == "code").font.name == "Cascadia Mono"


def test_unicode_round_trips(tmp_path: Path) -> None:
    output = render(tmp_path, "# Café\n\n日本語 and 😀 remain editable.\n")
    document = Document(output)
    assert document.paragraphs[0].text == "Café"
    assert document.paragraphs[1].text == "日本語 and 😀 remain editable."
