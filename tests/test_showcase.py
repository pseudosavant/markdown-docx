from __future__ import annotations

from pathlib import Path

from docx import Document

from markdown_docx.models import (
    CodeBlock,
    HeadingBlock,
    ImageBlock,
    ListParagraphBlock,
    PageBreakBlock,
    ParagraphBlock,
    SectionBreakBlock,
    TableBlock,
)
from markdown_docx.parser import parse_document
from markdown_docx.renderer import render_docx

PROJECT_ROOT = Path(__file__).resolve().parents[1]
SHOWCASE_PATH = PROJECT_ROOT / "sample" / "showcase.md"


def _showcase_model():
    source = SHOWCASE_PATH.read_text(encoding="utf-8")
    return source, parse_document(source, input_path=SHOWCASE_PATH, source_name=str(SHOWCASE_PATH))


def test_showcase_covers_portable_markdown_and_metadata_features() -> None:
    source, model = _showcase_model()

    assert model.options.section.page_size.name == "letter"
    assert model.options.section.orientation == "portrait"
    assert model.options.styles.paragraph == "Normal"
    assert model.options.styles.headings == {level: f"Heading {level}" for level in range(1, 7)}
    assert model.options.styles.blockquote == "Quote"
    assert model.options.styles.code_block == "Code Block"
    assert model.options.styles.ordered_list == ["List Number", "List Number 2", "List Number 3"]
    assert model.options.styles.unordered_list == ["List Bullet", "List Bullet 2", "List Bullet 3"]
    assert model.options.styles.table == "Table Grid"
    assert model.options.fonts.body == "Calibri"
    assert model.options.fonts.headings == "Calibri"
    assert model.options.fonts.monospace == "Consolas"

    block_types = {type(block) for block in model.blocks}
    assert block_types == {
        CodeBlock,
        HeadingBlock,
        ImageBlock,
        ListParagraphBlock,
        PageBreakBlock,
        ParagraphBlock,
        SectionBreakBlock,
        TableBlock,
    }

    headings = [block for block in model.blocks if isinstance(block, HeadingBlock)]
    assert {heading.level for heading in headings} == set(range(1, 7))

    text_fragments = [
        fragment
        for block in model.blocks
        if isinstance(block, (ParagraphBlock, HeadingBlock, ListParagraphBlock))
        for fragment in block.fragments
    ]
    assert any(fragment.bold and not fragment.italic for fragment in text_fragments)
    assert any(fragment.italic and not fragment.bold for fragment in text_fragments)
    assert any(fragment.bold and fragment.italic for fragment in text_fragments)
    assert any(fragment.code for fragment in text_fragments)
    assert any(fragment.kind == "break" for fragment in text_fragments)
    assert any(fragment.kind == "image" for fragment in text_fragments)
    assert any(
        "source line and continues through a soft source break" in (fragment.text or "") for fragment in text_fragments
    )

    list_blocks = [block for block in model.blocks if isinstance(block, ListParagraphBlock)]
    assert {(block.list_kind, block.depth) for block in list_blocks} == {
        ("ordered", 0),
        ("ordered", 1),
        ("ordered", 2),
        ("unordered", 0),
        ("unordered", 1),
        ("unordered", 2),
    }

    tables = [block for block in model.blocks if isinstance(block, TableBlock)]
    assert {table.options.alignment for table in tables} == {"left", "center", "right"}
    assert {table.options.width for table in tables} == {"auto", "page"}
    assert any(table.options.column_widths is None for table in tables)
    assert any(table.options.column_widths is not None for table in tables)
    assert {cell.alignment for table in tables for cell in table.headers} == {"left", "center", "right"}
    table_fragments = [
        fragment
        for table in tables
        for row in [table.headers, *table.rows]
        for cell in row
        for fragment in cell.fragments
    ]
    assert any(fragment.bold for fragment in table_fragments)
    assert any(fragment.italic for fragment in table_fragments)
    assert any(fragment.code for fragment in table_fragments)

    images = [block for block in model.blocks if isinstance(block, ImageBlock)]
    assert {image.options.alignment for image in images} == {"left", "center", "right"}
    assert any(image.options.width is None for image in images)
    assert any(image.options.width_is_percent for image in images)
    assert all(not image.src.startswith(("http://", "https://")) for image in images)
    for unit_example in ("92%", "5.9in", "15cm", "145mm", "360pt"):
        assert unit_example in source

    sections = [block.settings for block in model.blocks if isinstance(block, SectionBreakBlock)]
    assert [section.page_size.name for section in sections] == ["a4", "legal", None, "letter"]
    assert {section.orientation for section in sections} == {"portrait", "landscape"}
    assert len([block for block in model.blocks if isinstance(block, PageBreakBlock)]) >= 4


def test_showcase_renders_as_editable_native_word_content(tmp_path: Path) -> None:
    _, model = _showcase_model()
    output = tmp_path / "showcase.docx"

    result = render_docx(
        model,
        output,
        template_path=None,
        base_dir=SHOWCASE_PATH.parent,
        allow_remote_images=False,
    )

    document = Document(output)
    assert result["sections"] == 5
    assert result["warnings"] == ["image_alt_text_not_embedded"]
    assert len(document.sections) == 5
    assert len(document.tables) == 4
    assert len(document.inline_shapes) == 8
    assert any(paragraph.text == "The Great Lunch Bag Chase" for paragraph in document.paragraphs)
    assert any(paragraph.text == "Back to Document Defaults" for paragraph in document.paragraphs)
