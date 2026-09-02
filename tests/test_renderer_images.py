from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

from markdown_docx.errors import AssetError
from markdown_docx.parser import parse_document
from markdown_docx.renderer import render_docx


def test_standalone_image_width_alignment_and_warning(tmp_path: Path, png_file: Path) -> None:
    source = """<!-- markdown-docx
image:
  width: 50%
  alignment: center
-->

![Pixel](image.png)
"""
    model = parse_document(source, input_path=tmp_path / "input.md", source_name="input.md")
    output = tmp_path / "image.docx"
    result = render_docx(model, output, template_path=None, base_dir=tmp_path, allow_remote_images=False)
    document = Document(output)
    assert document.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert document.inline_shapes[0].width.inches == pytest.approx(3.25, abs=0.01)
    assert document.inline_shapes[0].height.inches == pytest.approx(3.25, abs=0.01)
    assert result["warnings"] == ["image_alt_text_not_embedded"]


def test_inline_image_stays_in_text_paragraph(tmp_path: Path, png_file: Path) -> None:
    source = "Before ![pixel](image.png) after.\n"
    model = parse_document(source, input_path=tmp_path / "input.md", source_name="input.md")
    output = tmp_path / "inline.docx"
    render_docx(model, output, template_path=None, base_dir=tmp_path, allow_remote_images=False)
    document = Document(output)
    assert len(document.paragraphs) == 1
    assert document.paragraphs[0].text == "Before  after."
    assert len(document.inline_shapes) == 1


def test_natural_image_is_clamped_to_usable_width(tmp_path: Path) -> None:
    image_path = tmp_path / "wide.png"
    Image.new("RGB", (1000, 100), "white").save(image_path, dpi=(10, 10))
    model = parse_document("![wide](wide.png)\n", input_path=tmp_path / "input.md", source_name="input.md")
    output = tmp_path / "wide.docx"
    render_docx(model, output, template_path=None, base_dir=tmp_path, allow_remote_images=False)
    shape = Document(output).inline_shapes[0]
    assert shape.width.inches == pytest.approx(6.5, abs=0.01)
    assert shape.height.inches == pytest.approx(0.65, abs=0.01)


def test_explicit_image_width_cannot_exceed_section(tmp_path: Path, png_file: Path) -> None:
    source = "<!-- markdown-docx\nimage:\n  width: 7in\n-->\n\n![pixel](image.png)\n"
    model = parse_document(source, input_path=tmp_path / "input.md", source_name="input.md")
    with pytest.raises(AssetError) as excinfo:
        render_docx(
            model,
            tmp_path / "too-wide.docx",
            template_path=None,
            base_dir=tmp_path,
            allow_remote_images=False,
        )
    assert excinfo.value.context.code == "image_too_wide"


def test_missing_and_remote_disabled_images_have_stable_errors(tmp_path: Path) -> None:
    missing = parse_document("![missing](none.png)\n", input_path=tmp_path / "input.md", source_name="input.md")
    with pytest.raises(AssetError) as missing_error:
        render_docx(
            missing, tmp_path / "missing.docx", template_path=None, base_dir=tmp_path, allow_remote_images=False
        )
    assert missing_error.value.context.code == "image_not_found"

    remote = parse_document(
        "![remote](https://example.com/image.png)\n",
        input_path=tmp_path / "input.md",
        source_name="input.md",
    )
    with pytest.raises(AssetError) as remote_error:
        render_docx(remote, tmp_path / "remote.docx", template_path=None, base_dir=tmp_path, allow_remote_images=False)
    assert remote_error.value.context.code == "image_download_failed"
