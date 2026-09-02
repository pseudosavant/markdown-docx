from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from markdown_docx.parser import parse_document
from markdown_docx.renderer import render_docx


def test_table_style_alignment_widths_and_cell_alignment(tmp_path: Path) -> None:
    source = """<!-- markdown-docx
table:
  style: Table Grid
  alignment: center
  width: page
  column_widths: [3, 1]
-->

| Item | Count |
| :--- | ---: |
| Widget | **2** |
"""
    model = parse_document(source, input_path=tmp_path / "input.md", source_name="input.md")
    output = tmp_path / "table.docx"
    render_docx(model, output, template_path=None, base_dir=tmp_path, allow_remote_images=False)
    table = Document(output).tables[0]
    assert table.style.name == "Table Grid"
    assert table.alignment == WD_TABLE_ALIGNMENT.CENTER
    assert table.columns[0].width.inches == pytest.approx(4.875, abs=0.02)
    assert table.columns[1].width.inches == pytest.approx(1.625, abs=0.02)
    assert table.cell(0, 0).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert table.cell(0, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert table.cell(1, 1).paragraphs[0].runs[0].bold is True


def test_table_uses_active_landscape_section_width(tmp_path: Path) -> None:
    source = """# Start

<!-- markdown-docx
section:
  orientation: landscape
  margins:
    left: 1in
    right: 1in
-->

| A | B |
| --- | --- |
| x | y |
"""
    model = parse_document(source, input_path=tmp_path / "input.md", source_name="input.md")
    output = tmp_path / "landscape-table.docx"
    render_docx(model, output, template_path=None, base_dir=tmp_path, allow_remote_images=False)
    table = Document(output).tables[0]
    assert sum(column.width.inches for column in table.columns) == pytest.approx(9, abs=0.02)
