from __future__ import annotations

import base64
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.text.paragraph import Paragraph

PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def package_parts(path: Path) -> set[str]:
    with ZipFile(path) as archive:
        return set(archive.namelist())


def test_python_docx_120_public_api_capability_matrix(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "output.docx"
    document = Document()
    document.save(source)
    source_parts = package_parts(source)

    paragraph_styles = [style for style in document.styles if style.type == WD_STYLE_TYPE.PARAGRAPH]
    table_styles = [style for style in document.styles if style.type == WD_STYLE_TYPE.TABLE]
    assert paragraph_styles
    assert table_styles
    assert document.styles["Normal"].type == WD_STYLE_TYPE.PARAGRAPH
    document.styles["Normal"].font.name = "Aptos"

    first_section = document.sections[0]
    first_section.top_margin = Inches(0.8)
    first_section.orientation = WD_ORIENT.PORTRAIT
    second_section = document.add_section(WD_SECTION.NEW_PAGE)
    second_section.orientation = WD_ORIENT.LANDSCAPE
    second_section.page_width = Inches(11)
    second_section.page_height = Inches(8.5)
    document.add_page_break()

    document.add_paragraph("Bullet", style="List Bullet")
    document.add_paragraph("Nested bullet", style="List Bullet 2")
    document.add_paragraph("Number", style="List Number")
    document.add_paragraph("Nested number", style="List Number 2")

    table = document.add_table(rows=2, cols=2, style="Table Grid")
    table.autofit = False
    for column, width in zip(table.columns, (Inches(4), Inches(2)), strict=True):
        column.width = width
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Count"

    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = image_paragraph.add_run().add_picture(BytesIO(PNG_1X1), width=Inches(1))
    assert shape.width == Inches(1)
    assert shape.height == Inches(1)

    document.save(output)
    reopened = Document(output)
    output_parts = package_parts(output)
    assert "word/theme/theme1.xml" in source_parts
    assert "word/numbering.xml" in source_parts
    assert source_parts <= output_parts
    assert len(reopened.sections) == 2
    assert reopened.styles["Normal"].font.name == "Aptos"
    assert len(reopened.tables) == 1
    assert len(reopened.inline_shapes) == 1


def test_unsupported_authoring_capabilities_are_not_public() -> None:
    assert not hasattr(Paragraph, "add_hyperlink")
    shape = Document().add_picture(BytesIO(PNG_1X1))
    assert not hasattr(shape, "alt_text")
    assert not hasattr(shape, "description")
