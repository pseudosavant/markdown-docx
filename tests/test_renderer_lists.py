from __future__ import annotations

from pathlib import Path

from docx import Document

from markdown_docx.parser import parse_document
from markdown_docx.renderer import render_docx


def test_ordered_unordered_nested_and_mixed_lists_use_real_styles(tmp_path: Path) -> None:
    source = """- Bullet
  1. Nested number
     - Deep bullet
- Second bullet

1. Number
   - Nested bullet
"""
    model = parse_document(source, input_path=tmp_path / "input.md", source_name="input.md")
    output = tmp_path / "lists.docx"
    render_docx(model, output, template_path=None, base_dir=tmp_path, allow_remote_images=False)
    document = Document(output)
    assert [paragraph.style.name for paragraph in document.paragraphs] == [
        "List Bullet",
        "List Number 2",
        "List Bullet 3",
        "List Bullet",
        "List Number",
        "List Bullet 2",
    ]
    assert [paragraph.text for paragraph in document.paragraphs] == [
        "Bullet",
        "Nested number",
        "Deep bullet",
        "Second bullet",
        "Number",
        "Nested bullet",
    ]
